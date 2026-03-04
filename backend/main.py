import asyncio
import base64
import io
import os
import time
from datetime import datetime
from typing import List, Optional

import vertexai
from vertexai.generative_models import GenerativeModel, Part
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from PIL import Image
import httpx
from pydantic import BaseModel
from supabase import Client, create_client
from docx import Document

load_dotenv()

# Vertex AI Configuration
PROJECT_ID = "stroyka-489218"
LOCATION = "us-central1"
EXAMINER_INSTRUCTIONS = "Structural engineer. Analyze defect in photo. Classify: A(исправное) B(работоспособное) C(ограниченно) D(аварийное). JSON format: {\"description\": \"...\", \"status_category\": \"A|B|C|D\"}"

# Initialize Vertex AI
vertexai.init(project=PROJECT_ID, location=LOCATION)

# Rate limiting
REQUEST_COUNT = 0
LAST_RESET = datetime.now()
DAILY_LIMIT = 100  # Maximum requests per day

def check_rate_limit():
    global REQUEST_COUNT, LAST_RESET
    
    now = datetime.now()
    if now.date() > LAST_RESET.date():
        REQUEST_COUNT = 0
        LAST_RESET = now
    
    if REQUEST_COUNT >= DAILY_LIMIT:
        raise HTTPException(
            status_code=429, 
            detail=f"Daily limit of {DAILY_LIMIT} requests exceeded. Try again tomorrow."
        )
    
    REQUEST_COUNT += 1
    print(f"DEBUG: Request count: {REQUEST_COUNT}/{DAILY_LIMIT}")

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_BUCKET_IMAGES = os.getenv("SUPABASE_BUCKET_IMAGES", "defect-images")
SUPABASE_BUCKET_REPORTS = os.getenv("SUPABASE_BUCKET_REPORTS", "defect-reports")
VERTEX_API_KEY = os.getenv("VERTEX_API_KEY")
EXAMINER_INSTRUCTIONS = os.getenv(
    "EXAMINER_INSTRUCTIONS",
    "You are a construction examiner following GOST standards. "
    "Given a structural photo, describe visible defects and assign a status category "
    "(A - no defect, B - minor, C - serious, D - critical).",
)

if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and GEMINI_API_KEY):
    # We don't raise immediately so the app can still start for local editing,
    # but requests that need these will fail with a clear error.
    print("WARNING: Backend missing SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY or GEMINI_API_KEY.")


supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)  # type: ignore[arg-type]

genai.configure(api_key=GEMINI_API_KEY)
# Используйте имя модели, доступное в вашем API (AI Studio или Vertex)
gemini_model = genai.GenerativeModel("gemini-1.5-flash-002")


class SurveyCreate(BaseModel):
    user_id: str
    name: str
    industry_gost: str


class Survey(BaseModel):
    id: str
    user_id: str
    name: str
    industry_gost: str


class DefectAnalyzeRequest(BaseModel):
    survey_id: str
    image_url: Optional[str] = None
    image_base64: Optional[str] = None
    axis: str  # "X" or "Y"
    construction_type: str  # "Concrete" | "Brick" | "Metal" | "Roof"
    location: Optional[str] = None


class Defect(BaseModel):
    id: str
    survey_id: str
    image_url: str
    axis: str
    construction_type: str
    location: Optional[str]
    description: str
    status_category: str


app = FastAPI(title="Industrial Defect Examiner API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/surveys", response_model=Survey)
async def create_survey(payload: SurveyCreate):
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(status_code=500, detail="Supabase is not configured.")

    data = {
        "user_id": payload.user_id,
        "name": payload.name,
        "industry_gost": payload.industry_gost,
    }
    try:
        res = supabase.table("surveys").insert(data).execute()
    except Exception as e:  # pragma: no cover - supabase client errors
        raise HTTPException(status_code=500, detail=f"Failed to create survey: {e}")

    if not res.data:
        raise HTTPException(status_code=500, detail="Supabase did not return survey.")

    row = res.data[0]
    return Survey(
        id=str(row["id"]),
        user_id=row["user_id"],
        name=row["name"],
        industry_gost=row["industry_gost"],
    )


@app.get("/surveys", response_model=List[Survey])
async def list_surveys(user_id: str):
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(status_code=500, detail="Supabase is not configured.")

    try:
        res = supabase.table("surveys").select("*").eq("user_id", user_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch surveys: {e}")

    return [
        Survey(
            id=str(row["id"]),
            user_id=row["user_id"],
            name=row["name"],
            industry_gost=row["industry_gost"],
        )
        for row in (res.data or [])
    ]


async def _fetch_image_as_pil(url: str) -> Image.Image:
    if not url:
        raise HTTPException(status_code=400, detail="image_url is required when image_base64 is not provided.")
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(url)
        if resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Could not fetch image: {resp.status_code}")
        return Image.open(io.BytesIO(resp.content)).convert("RGB")


def _build_system_prompt() -> str:
    return (
        f"{EXAMINER_INSTRUCTIONS}\n\n"
        "Return a concise JSON object with keys:\n"
        '"description": string (defect description using GOST terminology when possible),\n'
        '"status_category": string (one of: A, B, C, D).\n'
        "Respond ONLY with valid JSON."
    )


def _parse_gemini_response(raw: str) -> dict[str, str]:
    """Parse Gemini response to extract description and status_category."""
    try:
        # Try to parse as JSON first
        data = json.loads(raw)
        if isinstance(data, dict):
            return {
                "description": data.get("description", "No description provided."),
                "status_category": data.get("status_category", "B")
            }
    except json.JSONDecodeError:
        pass
    
    # If not JSON, try to extract from text
    description = "No description provided."
    status_category = "B"
    
    # Look for description patterns
    if "[ОПИСАНИЕ ТЕХНИЧЕСКОГО ДЕФЕКТА]" in raw:
        parts = raw.split("[ОПИСАНИЕ ТЕХНИЧЕСКОГО ДЕФЕКТА]")
        if len(parts) > 1:
            desc_part = parts[1].split("[")[0].strip()
            description = desc_part
    elif "description" in raw.lower():
        # Try to extract JSON-like description
        import re
        desc_match = re.search(r'description["\s]*:["\s]*"([^"]+)"', raw, re.IGNORECASE)
        if desc_match:
            description = desc_match.group(1)
    
    # Look for status category patterns
    if "[КАТЕГОРИЯ]" in raw:
        parts = raw.split("[КАТЕГОРИЯ]")
        if len(parts) > 1:
            status_part = parts[1].split("[")[0].strip()
            status_category = status_part[0] if status_part else "B"
    elif "status_category" in raw.lower():
        import re
        status_match = re.search(r'status_category["\s]*:["\s]*"([ABCD])"', raw, re.IGNORECASE)
        if status_match:
            status_category = status_match.group(1)
    
    return {
    image: Image.Image,
    axis: str,
    construction_type: str,
    location: Optional[str],
) -> dict:
    print(f"DEBUG: Vertex AI analysis started for {construction_type}, axis {axis}")
    
    try:
        from google.cloud import aiplatform
        from google.cloud.aiplatform import gapic
        from google.oauth2 import service_account
        
        # Initialize Vertex AI client
        credentials = service_account.Credentials.from_service_account_file(
            "path/to/service_account_key.json",
            scopes=["https://www.googleapis.com/auth/cloud-platform"],
        )
        aiplatform.init(project="your-project-id", credentials=credentials)
        
        # Create Vertex AI model resource
        model_resource = aiplatform.Model(
            display_name="your-model-display-name",
            location="your-model-location",
            project="your-project-id",
        )
        
        # Convert PIL image to bytes
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format="JPEG")
        img_bytes = img_byte_arr.getvalue()
        print(f"DEBUG: Image converted to bytes, size: {len(img_bytes)}")
        
        # Create image part for Vertex AI
        image_part = gapic.Image(image_bytes=img_bytes, mime_type="image/jpeg")
        print("DEBUG: Vertex AI image part created")
        
        prompt = f"""
{EXAMINER_INSTRUCTIONS}

Axis: {axis}
Type: {construction_type}
Location: {location or 'N/A'}

JSON format: {{"description": "...", "status_category": "A|B|C|D"}}
"""
        print(f"DEBUG: Prompt length: {len(prompt)} characters")
        
        # Generate content
        response = await asyncio.to_thread(
            model_resource.predict(
                instances=[{"prompt": prompt, "image": image_part}],
                parameters={"temperature": 0.1, "max_output_tokens": 500},
            )
        )
        print("DEBUG: Vertex AI response received")
        
        raw = response.predictions[0].text
        print(f"DEBUG: Vertex AI raw response: {raw[:200]}...")
        
        data = _parse_gemini_response(raw)
        print(f"DEBUG: Parsed response: {data}")
        return data
        
    except Exception as e:
        print(f"DEBUG: Vertex AI analysis failed: {e}")
        print(f"DEBUG: Exception type: {type(e)}")
        print(f"DEBUG: Exception args: {e.args}")
        return {
            "description": f"Analysis failed: {str(e)}",
            "status_category": "B"
        }

# ... (rest of the code remains the same)

@app.post("/defects/analyze", response_model=Defect)
async def analyze_defect(payload: DefectAnalyzeRequest):
    check_rate_limit()

    if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY and GEMINI_API_KEY):
        raise HTTPException(status_code=500, detail="Backend is missing environment configuration.")

    image_bytes: Optional[bytes] = None
    if payload.image_base64:
        try:
            image_bytes = base64.b64decode(payload.image_base64)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to decode base64 image: {e}")
        try:
            image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Uploaded image is not a valid image: {e}")
    else:
        image = await _fetch_image_as_pil(payload.image_url or "")

    gemini_result = await _analyze_defect_with_gemini(
        image=image,
        axis=payload.axis,
        construction_type=payload.construction_type,
        location=payload.location,
    )

    image_url_to_store = payload.image_url
    if image_bytes is not None:
        try:
            filename = f"survey-{payload.survey_id}-{int(time.time())}.jpg"
            supabase.storage.from_(SUPABASE_BUCKET_IMAGES).upload(
                path=filename,
                file=image_bytes,
                file_options={"content-type": "image/jpeg"},
            )
            image_url_to_store = supabase.storage.from_(SUPABASE_BUCKET_IMAGES).get_public_url(filename)
                
        except Exception as e:
            base = (SUPABASE_URL or "").rstrip("/")
            image_url_to_store = f"{base}/storage/v1/object/public/{SUPABASE_BUCKET_IMAGES}/{filename}" if base else f"placeholder-{filename}"

    defect_row = {
        "survey_id": payload.survey_id,
        "image_url": image_url_to_store,
        "axis": payload.axis,
        "construction_type": payload.construction_type,
        "location": payload.location,
        "description": gemini_result["description"],
        "status_category": gemini_result["status_category"],
    }

    try:
        res = supabase.table("defects").insert(defect_row).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save defect: {e}")

    if not res.data:
        raise HTTPException(status_code=500, detail="Supabase did not return defect.")

    row = res.data[0]
    return Defect(
        id=str(row["id"]),
        survey_id=row["survey_id"],
        image_url=row["image_url"],
        axis=row["axis"],
        construction_type=row["construction_type"],
        location=row.get("location"),
        description=row["description"],
        status_category=row["status_category"],
    )


@app.get("/defects/{survey_id}", response_model=List[Defect])
async def get_defects_for_survey(survey_id: str):
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(status_code=500, detail="Supabase is not configured.")

    try:
        res = supabase.table("defects").select("*").eq("survey_id", survey_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch defects: {e}")

    return [
        Defect(
            id=str(row["id"]),
            survey_id=row["survey_id"],
            image_url=row["image_url"],
            axis=row["axis"],
            construction_type=row["construction_type"],
            location=row.get("location"),
            description=row["description"],
            status_category=row["status_category"],
        )
        for row in (res.data or [])
    ]


@app.get("/reports/{survey_id}")
async def generate_report(survey_id: str):
    if not (SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY):
        raise HTTPException(status_code=500, detail="Supabase is not configured.")

    try:
        defects_res = supabase.table("defects").select("*").eq("survey_id", survey_id).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch defects: {e}")

    defects = defects_res.data or []

    if not defects:
        raise HTTPException(status_code=404, detail="No defects found for survey.")

    # Build DOCX in memory
    doc = Document()
    doc.add_heading("Construction Defect Report", level=1)
    doc.add_paragraph(f"Survey ID: {survey_id}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Photo"
    hdr_cells[1].text = "Location"
    hdr_cells[2].text = "Description of the defect"
    hdr_cells[3].text = "Status category"

    for d in defects:
        row_cells = table.add_row().cells
        row_cells[0].text = d.get("image_url", "")
        row_cells[1].text = d.get("location") or ""
        row_cells[2].text = d.get("description") or ""
        row_cells[3].text = d.get("status_category") or ""

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"survey-{survey_id}-report.docx"

    # Optionally upload to Supabase storage
    try:
        supabase.storage.from_(SUPABASE_BUCKET_REPORTS).upload(
            file=buffer.getvalue(),
            path=filename,
            file_options={"content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        )
    except Exception as e:
        # Non-fatal: report is still streamed back to client
        print(f"WARNING: Failed to upload report to Supabase: {e}")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# To run locally:
# uvicorn main:app --reload --host 0.0.0.0 --port 8000

